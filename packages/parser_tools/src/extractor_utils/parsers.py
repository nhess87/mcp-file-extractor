from io import BytesIO
import re
import fitz
import io
from PIL import Image, ImageFilter
from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
#from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence.models import DocumentContentFormat
import os
import google.generativeai as genai
import base64
from openai import OpenAI
from openai import APIConnectionError, RateLimitError, APIError
import base64
from azure.storage.blob import BlobServiceClient
import imghdr
import logging


logger = logging.getLogger("tasks")


def extract_text_from_image_with_azure(image_data: str):
    
    extracted_text=""
    client = DocumentIntelligenceClient(
        os.getenv('AZURE_DOC_INTEL_ENDPOINT'),
        AzureKeyCredential(os.getenv('AZURE_DOC_INTEL_KEY'))
    )

    image_bytes= BytesIO(base64.b64decode(image_data))
    
    detected_format= imghdr.what(image_bytes)
    content_type = f"image/{detected_format}" if detected_format else "image/png"
    
    file_data = image_bytes.getvalue()
    

    # Check for empty file data
    if not file_data:
        raise ValueError("Uploaded file is empty")

    try:
        poller = client.begin_analyze_document(
            "prebuilt-layout",
            file_data,
            content_type=content_type,
            output_content_format="markdown" 
        )
        result = poller.result()

        extracted_text = result.content
        return extracted_text
    except Exception as e:
        logger.info(f"Error analyzing image: {e}")
        return ""

def extract_text_from_image_with_gemini_multimodal(image_data: str):
    import os
    import base64
    from io import BytesIO
    import tiktoken
    import imghdr

    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
    extracted_text=""
    try:
        image_bytes = BytesIO(base64.b64decode(image_data))
        detected_format = imghdr.what(image_bytes)  # Returns 'jpeg', 'png', etc.
        if detected_format:
            mime_type = f"image/{detected_format}"  # Convert to proper MIME type format
        else:
            mime_type = "image/png"
        image_bytes.seek(0)
        # Initialize Gemini model
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(
            [
                "Extract all text from this image accurately. Return only the raw extracted text.",
                {
                    "mime_type": mime_type,
                    "data": image_bytes.read()
                }
            ]
        )
        extracted_text += response.text.strip() if response.text else "No text found"
        
        return extracted_text
    except Exception as e:
        return f"An error occurred: {str(e)}"


def extract_text_from_pdf_with_azure(file_data: str):
    # Initialize Azure Document Analysis client
    ##client = DocumentAnalysisClient(endpoint, AzureKeyCredential(key)) --old version
    print(f"this function is triggered")
    logger.info(f"this is azure doc endpoint: {os.getenv('AZURE_DOC_INTEL_ENDPOINT')}")
    client =  DocumentIntelligenceClient(os.getenv('AZURE_DOC_INTEL_ENDPOINT'), AzureKeyCredential(os.getenv('AZURE_DOC_INTEL_KEY')))
    logger.info(f"client: {client}")
    file_bytes= BytesIO(base64.b64decode(file_data))
    file_data= file_bytes.getvalue()
    logger.info(f"this is the file data: {file_data}")
    print(f"this is the file data: {file_data}")

    if not file_data:
        raise ValueError("Uploaded file is empty")
    try:
        
        # Use Azure service for document analysis
        poller = client.begin_analyze_document("prebuilt-layout", file_data, output_content_format=DocumentContentFormat.MARKDOWN)
        result = poller.result()
        # Extract text from Azure results
        extracted_text = result.content
    
        """
        for page in result.pages:
            for line in page.lines:
                extracted_text += line.content + "\n" 

            for table in result.tables:
                for cell in table.cells:
                    extracted_text += cell.content + "\t"
                extracted_text += "\n"
        """
        return extracted_text
    except Exception as e:
        logger.info(f"Error analyzing image: {e}")
        return ""


def extract_text_from_pdf_with_gemini_multimodal(file_data: str):
    import os
    import base64
    from io import BytesIO
    import imghdr
    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
    try:
        file_bytes= BytesIO(base64.b64decode(file_data))

        mime_type = "application/pdf"  
        # Read the file content
        pdf_bytes = file_bytes.getvalue()
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(
            [
                "Extract all text from this PDF document accurately. Return only the raw extracted text.",
                {
                    "mime_type": mime_type,
                    "data": pdf_bytes
                }
            ]
        )
        extracted_text= response.text.strip() if response.text else "No text found"

        return extracted_text
    except Exception as e:
        return f"An error occurred: {str(e)}"

def extract_text_from_image_with_openai_multimodal(image_data: str):
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    extracted_text=""
    try:
        # Decode base64 string into bytes
        image_bytes = base64.b64decode(image_data)

        detected_format= imghdr.what(None,image_bytes)
        mime_type = f"image/{detected_format}" if detected_format else "image/png"
        image_data= f"data:{mime_type};base64,{image_data}"

        # Create API request
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Extract all text from this image exactly as it appears. Return only the raw text without any formatting or commentary."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": image_data
                            }
                        }
                    ]
                }
            ],
            max_tokens=1000
        )

        if response.choices and response.choices[0].message.content:
            extracted_text += response.choices[0].message.content.strip()

        return extracted_text
    except Exception as e:
        return f"Error: {str(e)}"


def extract_text_from_pdf_with_openai_multimodal(file_data:str):
    extracted_text=""
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    try:
        file_bytes= BytesIO(base64.b64decode(file_data))
        pdf_bytes= file_bytes.getvalue()
        
        # Convert PDF to images
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        full_text = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(dpi=200)  # Render page as image
            img_bytes = BytesIO(pix.tobytes("PNG"))
            
            # Convert to base64
            base64_image = base64.b64encode(img_bytes.getvalue()).decode("utf-8")
            image_data = f"data:image/png;base64,{base64_image}"
            
            # Send to OpenAI
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Extract all text exactly as shown. Preserve formatting, tables, and layout."},
                            {"type": "image_url", "image_url": {"url": image_data}}
                        ]
                    }
                ],
                max_tokens=2000
            )
            
            if response.choices and response.choices[0].message.content:
                full_text.append(response.choices[0].message.content.strip())
                
        extracted_text += "\n\n".join(full_text) if full_text else "No text found"

        return extracted_text
    except Exception as e:
        return f"Error: {str(e)}"

def extract_text_from_txt(uploaded_file):
    try:
        return uploaded_file.read().decode('utf-8')
    except Exception as e:
        print(f"Error reading TXT file: {e}")
        return ""

def extract_text_from_word_with_customfunction(id, file_data:str,storageConnString):
    import time
    text = ""  
    try:
        file_bytes= BytesIO(base64.b64decode(file_data))
        doc = DocxDocument(file_bytes)

        for para in doc.paragraphs:
            para_text = []
            for run in para.runs:
                hyperlink = run._element.getparent().xpath('.//w:hyperlink')
                if hyperlink:
                    rel_id = hyperlink[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')

                    if rel_id and rel_id in doc.part.rels:
                        # Get the actual URL
                        url = doc.part.rels[rel_id]._target  
                        display_text = hyperlink[0].text        
                        # Add the display text and URL to para_text
                        para_text.append(f"{display_text} ({url})")  
                        continue  
                # If not a hyperlink, just add the text as is
                para_text.append(run.text)
            text += ''.join(para_text) + "\n"
        image_counter = 0
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image = doc.part.rels[rel].target_part.blob
                img = Image.open(io.BytesIO(image))

                storage_conn_string = storageConnString
                blob_service_client = BlobServiceClient.from_connection_string(storage_conn_string)
                blob_name = f"{id}_{rel}_{image_counter}_{int(time.time() * 1000)}.png"
                blob_client = blob_service_client.get_blob_client(container="tem-folder", blob=blob_name)
                img_bytes = io.BytesIO()
                img.save(img_bytes, format='PNG')
                img_bytes.seek(0)
                blob_client.upload_blob(img_bytes, overwrite=True)

                downloaded_blob = blob_client.download_blob()
                image_data = downloaded_blob.readall()
                image_stream = io.BytesIO(image_data)
                image_stream.seek(0)

                
                extracted_text = extract_text_from_image_with_gemini_multimodal_reg(image_stream,"image/png")
                text += "\n" + extracted_text
                image_counter += 1
        
        return text
    except Exception as e:
        logger.info(f"Error reading document: {e}")
        return ""

def extract_text_from_image_with_gemini_multimodal_reg(image_bytes: BytesIO, mime_type="image/png"):
    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
    try:
        image_bytes.seek(0)
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(
            [
                "Extract all text from this image accurately. Return only the raw extracted text.",
                {
                    "mime_type": mime_type, 
                    "data": image_bytes.read()
                }
            ]
        )
        return response.text.strip() if response.text else "No text found"
    except Exception as e:
        return f"An error occurred: {str(e)}"
    

def extract_only_text_from_pdf(file_data: str):
    import PyPDF2 
    try:
        text = ""
        file_bytes= BytesIO(base64.b64decode(file_data))
        pdf_reader = PyPDF2.PdfReader(file_bytes)
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text: 
                text += page_text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error reading PDF file: {e}")
        return ""

def extract_only_text_from_word(file_data: str):
    import base64
    import io
    from io import BytesIO
    from spire.doc import Document as SpireDocument, DocumentObjectType, FieldType
    from docx import Document as DocxDocument

    try:
        # Decode base64 string to binary
        file_bytes = BytesIO(base64.b64decode(file_data))
        file_bytes.seek(0)  # Ensure the stream is at the beginning
        print("File successfully decoded into BytesIO.")

        text=""
        file_bytes.seek(0)  # Reset file pointer before retrying
        doc = DocxDocument(file_bytes)

        for para in doc.paragraphs:
            para_text = []
            for run in para.runs:
                hyperlink = run._element.getparent().xpath('.//w:hyperlink')
                if hyperlink:
                    rel_id = hyperlink[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')

                    if rel_id and rel_id in doc.part.rels:
                        # Get the actual URL
                        url = doc.part.rels[rel_id]._target  
                        display_text = hyperlink[0].text        
                        # Add the display text and URL to para_text
                        para_text.append(f"{display_text} ({url})")  
                        continue  
                # If not a hyperlink, just add the text as is
                para_text.append(run.text)
            text += ''.join(para_text) + "\n"
        return text
    except Exception as e2:
        print(f"Error using python-docx: {e2}")

    except Exception as e:
        print(f"Critical error processing Word file: {e}")

    return ""

def extract_text_function(extractorOCR,StorageConnString, blob_name, uploaded_file):
    try:

        logger.info(f"ocr option choosen: {extractorOCR}")
        logger.info(f"uploaded file is: {blob_name}")

        file_ext = blob_name.lower()

        if file_ext.lower().endswith(('.jpg', '.jpeg', '.png')):
            ocr_extractors = {
                "azure_documentIntelligence_layout": extract_text_from_image_with_azure,
                "gemini_multimodal": extract_text_from_image_with_gemini_multimodal,
                "openai_multimodal": extract_text_from_image_with_openai_multimodal,
            }
            ocr_function = ocr_extractors.get(extractorOCR)
            if not ocr_function:
                return "Error: Please provide a valid OCR extractor to process image files."

            uploaded_file.seek(0)
            file_data = base64.b64encode(uploaded_file.read()).decode("utf-8")
            extracted_text = ocr_function(file_data)
            return extracted_text
        
        ###PDF handling
        elif file_ext.endswith('.pdf'):
            ocr_extractors = {
                "azure_documentIntelligence_layout": extract_text_from_pdf_with_azure,
                "gemini_multimodal": extract_text_from_pdf_with_gemini_multimodal,
                "openai_multimodal": extract_text_from_pdf_with_openai_multimodal,
            }
            ocr_function = ocr_extractors.get(extractorOCR)
            
            uploaded_file.seek(0)
            file_data = base64.b64encode(uploaded_file.read()).decode("utf-8")

            if not ocr_function:
                extracted_text = extract_only_text_from_pdf(file_data)
                return extracted_text
            extracted_text = ocr_function(file_data)
            return extracted_text
        
        ##WORD DOC HANDLING
        elif file_ext.endswith(('.docx', '.doc')):
            uploaded_file.seek(0)
            file_data = base64.b64encode(uploaded_file.read()).decode("utf-8")
            id= 0

            if not extractorOCR:
                extracted_text = extract_only_text_from_word(file_data)
                return extracted_text
            extracted_text = extract_text_from_word_with_customfunction(id, file_data,StorageConnString)
            return extracted_text
        
        # TXT HANDLING
        elif file_ext.endswith('.txt'):
            uploaded_file.seek(0)
            extracted_text = extract_text_from_txt(uploaded_file)
            if not extracted_text:
                return "Error: No text found in TXT."
            return f"Uploaded Document Content: {extracted_text}"

        else:
            return f"Error: Unsupported file type {file_ext}"

    except Exception as e:
        logger.error(f"Exception occurred: {str(e)}")
        return f"Error in extracting text of the file: {str(e)}"